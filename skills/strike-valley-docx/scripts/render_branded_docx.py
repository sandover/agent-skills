#!/usr/bin/env python3
"""
Strike Valley Studio — Branded DOCX Renderer

Converts any markdown into a branded, professionally formatted .docx document.
Brand identity (fonts, colors, logo) is built in. Document metadata (title,
recipient, author, date) comes from command-line arguments.

Usage:
    python render_branded_docx.py input.md -o output.docx \\
        --title "Report Title" \\
        --prepared-for "Megan Harvey, Chief Revenue Officer" \\
        --author "Brandon Harvey, Strike Valley Studio" \\
        --date "26 Mar 2026" \\
        --toc after-first-section
"""
from __future__ import annotations

import argparse
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Emu, Inches, Pt, RGBColor

try:
    from PIL import Image, ImageDraw

    HAS_PIL = True
except ImportError:
    HAS_PIL = False


# ═══════════════════════════════════════════════════════════════════════════════
# Brand identity — Strike Valley Studio
# ═══════════════════════════════════════════════════════════════════════════════
EUROSTILE = "Eurostile"
EUROSTILE_BOLD = "Eurostile Bold"
IBM = "IBM Plex Sans"
INCONSOLATA = "Inconsolata"

DARK = "142134"  # Brand navy (confident-navy)
BLUE = "1B5E89"  # Brand blue
MID = "3A3632"  # Warm charcoal
H4COL = "5A7D8F"  # Desaturated blue
GRAY = "8A8A8A"  # Muted
BG = "FAF8F2"  # Warm off-white page
RULE_COL = DARK  # Match the deep navy logo
LINK_COL = DARK  # Keep URL text in the same deep navy family
WATERMARK_COL = "DDD9D0"  # Barely visible on FAF8F2

# Heading hierarchy: (font, size, color, bold, indent_in, space_before, space_after)
H_STYLES = {
    1: (EUROSTILE_BOLD, 22, DARK, True, 0, 36, 10),
    2: (EUROSTILE_BOLD, 15, BLUE, True, 0.33, 22, 6),
    3: (EUROSTILE, 13, MID, False, 0.66, 16, 4),
    4: (EUROSTILE, 12, H4COL, False, 1.0, 12, 2),
}

# Studio branding (constant across all documents)
STUDIO_LINES = ["STRIKE VALLEY", "STUDIO"]
STUDIO_URL = "https://strikevalley.studio"
STUDIO_URL_TEXT = "strikevalley.studio"
STUDIO_LOCATION = "Los Angeles, California"

# Assets — relative to skill root (scripts/../)
SKILL_ROOT = Path(__file__).resolve().parent.parent
DEEP_NAVY_LOGO_PATH = SKILL_ROOT / "assets" / "logo-blue-variants" / "oklch-round2" / "sun-solid@512-confident-navy-142134.png"
LOGO_PATH = DEEP_NAVY_LOGO_PATH if DEEP_NAVY_LOGO_PATH.exists() else (SKILL_ROOT / "assets" / "sun-solid@512.png")
EUROSTILE_REGULAR_FONT_PATH = Path("/Users/brandonharvey/Library/Fonts/Eurostile.ttf")


# ═══════════════════════════════════════════════════════════════════════════════
# Low-level helpers
# ═══════════════════════════════════════════════════════════════════════════════
def hex_rgb(v: str) -> RGBColor:
    return RGBColor.from_string(v.strip().lstrip("#").upper())


def set_font(
    run, font: str, size_pt: float, color_hex: str,
    bold: bool = False, italic: bool = False,
) -> None:
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), font)
    fonts.set(qn("w:hAnsi"), font)
    fonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size_pt)
    run.font.color.rgb = hex_rgb(color_hex)
    run.bold = bold
    run.italic = italic


def set_page_background(doc, color_hex: str) -> None:
    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), color_hex.strip().lstrip("#").upper())
    doc._element.insert(0, bg)


def clear_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "nil")


def clear_container(container) -> None:
    """Remove all existing block children from a header/footer container."""
    element = container._element
    for child in list(element):
        element.remove(child)


def set_table_fixed_width(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_cell_margins(cell, *, top=0, start=0, bottom=0, end=0) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, edge: str, spec: dict) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    el = tc_borders.find(qn(f"w:{edge}"))
    if el is None:
        el = OxmlElement(f"w:{edge}")
        tc_borders.append(el)
    for k, v in spec.items():
        el.set(qn(f"w:{k}"), v)


def set_runless_paragraph(paragraph) -> None:
    paragraph.text = ""
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0


def add_hyperlink(paragraph, *, text, url, font_name, size_pt, color_hex) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), font_name)
    rf.set(qn("w:hAnsi"), font_name)
    rf.set(qn("w:eastAsia"), font_name)
    rpr.append(rf)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color_hex.strip().lstrip("#").upper())
    rpr.append(c)
    s = OxmlElement("w:sz")
    s.set(qn("w:val"), str(int(size_pt * 2)))
    rpr.append(s)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hl.append(run)
    paragraph._p.append(hl)


# ---------------------------------------------------------------------------
# Field helpers (PAGE, PAGEREF)
# ---------------------------------------------------------------------------
def _make_field_rpr(font: str, size_pt: float, color_hex: str, position_half_pts: int | None = None):
    col = color_hex.strip().lstrip("#").upper()
    rpr = OxmlElement("w:rPr")
    f = OxmlElement("w:rFonts")
    f.set(qn("w:ascii"), font)
    f.set(qn("w:hAnsi"), font)
    f.set(qn("w:eastAsia"), font)
    rpr.append(f)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), col)
    rpr.append(c)
    s = OxmlElement("w:sz")
    s.set(qn("w:val"), str(int(size_pt * 2)))
    rpr.append(s)
    if position_half_pts is not None:
        pos = OxmlElement("w:position")
        pos.set(qn("w:val"), str(position_half_pts))
        rpr.append(pos)
    return rpr


def _field_run(*children, font, size_pt, color_hex, position_half_pts: int | None = None):
    r = OxmlElement("w:r")
    r.append(_make_field_rpr(font, size_pt, color_hex, position_half_pts=position_half_pts))
    for ch in children:
        r.append(ch)
    return r


def add_page_number_field(paragraph, font: str, size_pt: float, color_hex: str) -> None:
    """Insert a PAGE field — each component in its own run with full rPr."""
    kw = dict(font=font, size_pt=size_pt, color_hex=color_hex)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    p = paragraph._p
    p.append(_field_run(begin, **kw))
    p.append(_field_run(instr, **kw))
    p.append(_field_run(sep, **kw))
    p.append(_field_run(txt, **kw))
    p.append(_field_run(end, **kw))


def add_pageref_field(
    paragraph, bookmark: str, font: str, size_pt: float, color_hex: str,
) -> None:
    """Insert a PAGEREF field that resolves to the bookmark's page number."""
    kw = dict(font=font, size_pt=size_pt, color_hex=color_hex)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")  # auto-update on document open
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" PAGEREF {bookmark} \\h "
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "\u2013"  # placeholder until field update
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    p = paragraph._p
    p.append(_field_run(begin, **kw))
    p.append(_field_run(instr, **kw))
    p.append(_field_run(sep, **kw))
    p.append(_field_run(txt, **kw))
    p.append(_field_run(end, **kw))


# ---------------------------------------------------------------------------
# Bookmarks
# ---------------------------------------------------------------------------
_bm_id_counter = [0]


def add_bookmark(paragraph, name: str) -> None:
    bm_id = str(_bm_id_counter[0])
    _bm_id_counter[0] += 1
    bs = OxmlElement("w:bookmarkStart")
    bs.set(qn("w:id"), bm_id)
    bs.set(qn("w:name"), name)
    be = OxmlElement("w:bookmarkEnd")
    be.set(qn("w:id"), bm_id)
    paragraph._p.insert(0, bs)
    paragraph._p.append(be)


# ═══════════════════════════════════════════════════════════════════════════════
# Inline markdown parsing (**bold**, *italic*)
# ═══════════════════════════════════════════════════════════════════════════════
_INLINE_RE = re.compile(r"(\*\*.*?\*\*|\*[^*]+?\*)")


def add_rich_text(
    paragraph, text: str, font: str, size_pt: float, color_hex: str,
) -> None:
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            set_font(r, font, size_pt, color_hex, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1])
            set_font(r, font, size_pt, color_hex, italic=True)
        else:
            r = paragraph.add_run(part)
            set_font(r, font, size_pt, color_hex)


# ═══════════════════════════════════════════════════════════════════════════════
# Document-level setup
# ═══════════════════════════════════════════════════════════════════════════════
def setup_styles(doc) -> None:
    """Configure heading and Normal styles for brand identity + TOC support."""
    for level, (font, size, color, bold, indent, sb, sa) in H_STYLES.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = font
        style.font.size = Pt(size)
        style.font.color.rgb = hex_rgb(color)
        style.font.bold = bold
        pf = style.paragraph_format
        pf.space_before = Pt(sb)
        pf.space_after = Pt(sa)
        pf.left_indent = Inches(indent)
        pf.keep_with_next = True
        # eastAsia font
        rpr = style._element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rpr.insert(0, rf)
        rf.set(qn("w:ascii"), font)
        rf.set(qn("w:hAnsi"), font)
        rf.set(qn("w:eastAsia"), font)
        # Outline level for field-based TOC compatibility
        ppPr = style._element.find(qn("w:pPr"))
        if ppPr is None:
            ppPr = OxmlElement("w:pPr")
            style._element.append(ppPr)
        outline = ppPr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            ppPr.append(outline)
        outline.set(qn("w:val"), str(level - 1))

    # Normal style
    normal = doc.styles["Normal"]
    normal.font.name = IBM
    normal.font.size = Pt(11)
    normal.font.color.rgb = hex_rgb(DARK)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25
    rpr = normal._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    rf.set(qn("w:ascii"), IBM)
    rf.set(qn("w:hAnsi"), IBM)
    rf.set(qn("w:eastAsia"), IBM)


def create_watermark_logo(src: Path, dst: Path) -> Path | None:
    """Create a very faint version of the standalone logo for diagnostics."""
    if not HAS_PIL or not src.exists():
        return None
    try:
        img = Image.open(src).convert("RGBA")
        r, g, b, a = img.split()
        a = a.point(lambda x: int(x * 0.07))
        img = Image.merge("RGBA", (r, g, b, a))
        bg = Image.new("RGBA", img.size, (250, 248, 242, 255))
        result = Image.alpha_composite(bg, img).convert("RGB")
        result.save(str(dst), "PNG")
        return dst
    except Exception:
        return None


def create_footer_watermark(src: Path, dst: Path, *, opacity: float = 0.22) -> Path | None:
    """Create one combined footer watermark asset so logo and text are perfectly aligned."""
    if not HAS_PIL or not src.exists():
        return None
    try:
        logo = Image.open(src).convert("RGBA")
        logo.thumbnail((30, 30), Image.Resampling.LANCZOS)

        canvas_w, canvas_h = 680, 36
        canvas = Image.new("RGBA", (canvas_w, canvas_h), (0, 0, 0, 0))

        alpha_scale = max(0.0, min(opacity, 1.0))
        r, g, b, a = logo.split()
        a = a.point(lambda x: int(x * alpha_scale))
        logo = Image.merge("RGBA", (r, g, b, a))

        center_y = canvas_h / 2
        logo_y = round(center_y - (logo.height / 2))
        canvas.alpha_composite(logo, (0, logo_y))

        draw = ImageDraw.Draw(canvas)
        font_path = EUROSTILE_REGULAR_FONT_PATH if EUROSTILE_REGULAR_FONT_PATH.exists() else None
        if font_path:
            from PIL import ImageFont
            font = ImageFont.truetype(str(font_path), 22)
        else:
            from PIL import ImageFont
            font = ImageFont.load_default()
        text = "STRIKE VALLEY STUDIO"
        text_x = logo.width + 8
        text_bbox = draw.textbbox((0, 0), text, font=font)
        text_center_y = (text_bbox[1] + text_bbox[3]) / 2
        text_y = round(center_y - text_center_y)
        alpha = int(255 * alpha_scale)
        color = (0x15, 0x20, 0x28, alpha)
        draw.text((text_x, text_y), text, font=font, fill=color)
        canvas.save(str(dst), "PNG")
        return dst
    except Exception:
        return None


def create_header_rule_image(dst: Path, color_hex: str = RULE_COL) -> Path | None:
    """Create the branded top rule with the short right descender."""
    if not HAS_PIL:
        return None
    try:
        width, height = 1800, 68
        line_y = 4
        descender_bottom = 52
        stroke = 4
        color = tuple(int(color_hex[i:i + 2], 16) for i in (0, 2, 4)) + (255,)

        img = Image.new("RGBA", (width, height), (0, 0, 0, 0))
        draw = ImageDraw.Draw(img)
        x_end = width - 16
        draw.line((0, line_y, x_end, line_y), fill=color, width=stroke)
        draw.line((x_end, line_y, x_end, descender_bottom), fill=color, width=stroke)
        img.save(str(dst), "PNG")
        return dst
    except Exception:
        return None


def setup_page(doc, watermark_logo: Path | None = None) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    sec.footer_distance = Inches(0.3)
    set_page_background(doc, BG)

    # Auto-update fields on open (populates TOC page numbers)
    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    # Different first-page footer (no watermark on cover)
    sec.different_first_page_header_footer = True

    # First page keeps the same geometry as later pages, just without the watermark.
    build_footer_line(sec.first_page_footer, watermark_logo=watermark_logo, include_watermark=False)
    build_footer_line(sec.footer, watermark_logo=watermark_logo, include_watermark=True)


def build_footer_line(footer, *, watermark_logo: Path | None, include_watermark: bool) -> None:
    """Build a single-line footer with a left watermark image and right page number."""
    clear_container(footer)
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.92), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.SPACES
    )

    if include_watermark:
        if watermark_logo and watermark_logo.exists():
            p.add_run().add_picture(str(watermark_logo), width=Inches(2.0))
        else:
            text_run = p.add_run("STRIKE VALLEY STUDIO")
            set_font(text_run, EUROSTILE, 7.25, WATERMARK_COL)

    p.add_run("\t")
    add_page_number_field(p, font=EUROSTILE, size_pt=8.75, color_hex=GRAY)


# ═══════════════════════════════════════════════════════════════════════════════
# Branded header (top-matter / cover block)
# ═══════════════════════════════════════════════════════════════════════════════
def write_branded_header(doc, meta: dict) -> None:
    """Render the branded 2-column header: logo/brand left, title/metadata right.

    meta keys: title, prepared_for, author, date, version (all optional strings)
    """
    grid = doc.add_table(rows=1, cols=2)
    grid.alignment = WD_TABLE_ALIGNMENT.LEFT
    grid.autofit = False
    clear_table_borders(grid)

    left_w, right_w = 2.0, 5.0  # total 7.0" = page width minus margins
    grid.columns[0].width = Inches(left_w)
    grid.columns[1].width = Inches(right_w)
    row = grid.rows[0]
    row.cells[0].width = Inches(left_w)
    row.cells[1].width = Inches(right_w)
    row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    row.height = Inches(2.2)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    set_cell_margins(row.cells[0], top=24, start=24, bottom=120, end=30)
    set_cell_margins(row.cells[1], top=0, start=620, bottom=120, end=0)

    # --- Left cell: brand block (always the same) ---
    left = row.cells[0]
    spacer = left.paragraphs[0]
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(26)
    if LOGO_PATH.exists():
        lp = left.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run()
        lr.add_picture(str(LOGO_PATH), width=Inches(0.85))
        lp.paragraph_format.space_after = Pt(6)

    sp = left.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_after = Pt(10)
    sp.paragraph_format.line_spacing = 1.0
    for i, line in enumerate(STUDIO_LINES):
        if i:
            sp.add_run().add_break()
        r = sp.add_run(line)
        set_font(r, EUROSTILE_BOLD, 12, "111111", bold=True)

    wp = left.add_paragraph()
    wp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    wp.paragraph_format.space_after = Pt(1)
    wp.paragraph_format.space_before = Pt(0)
    add_hyperlink(
        wp,
        text=STUDIO_URL_TEXT,
        url=STUDIO_URL,
        font_name=INCONSOLATA,
        size_pt=9,
        color_hex=LINK_COL,
    )

    locp = left.add_paragraph()
    locp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    locp.paragraph_format.space_before = Pt(0)
    r = locp.add_run(STUDIO_LOCATION)
    set_font(r, IBM, 9, GRAY)

    # --- Right cell: title + metadata ---
    right = row.cells[1]
    rule_path = meta.get("header_rule_path")
    rule_p = right.paragraphs[0]
    rule_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rule_p.paragraph_format.space_before = Pt(0)
    rule_p.paragraph_format.space_after = Pt(0)
    if rule_path and Path(rule_path).exists():
        rule_run = rule_p.add_run()
        rule_run.add_picture(str(rule_path), width=Inches(4.74))
    title = meta.get("title", "")

    tp = right.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after = Pt(34)
    tp.paragraph_format.line_spacing = 1.0

    if title:
        # Split on ": " for natural multi-line titles
        if ": " in title:
            parts = title.split(": ", 1)
            r1 = tp.add_run(parts[0] + ":")
            set_font(r1, EUROSTILE, 22, "111111")
            tp.add_run().add_break()
            r2 = tp.add_run(parts[1])
            set_font(r2, EUROSTILE, 22, "111111")
        else:
            r = tp.add_run(title)
            set_font(r, EUROSTILE, 22, "111111")

    # Prepared-for block
    prepared_for = meta.get("prepared_for")
    if prepared_for:
        pp = right.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pp.paragraph_format.space_before = Pt(0)
        pp.paragraph_format.space_after = Pt(18)
        pp.paragraph_format.line_spacing = 1.15
        # Label
        lr = pp.add_run("prepared for:")
        set_font(lr, EUROSTILE, 10, GRAY)
        pp.add_run().add_break()
        # Parse "Name, Role" if comma present
        if ", " in prepared_for:
            name, role = prepared_for.split(", ", 1)
            nr = pp.add_run(name)
            set_font(nr, IBM, 13, "111111", bold=True)
            pp.add_run().add_break()
            rr = pp.add_run(role)
            set_font(rr, IBM, 11, "111111")
        else:
            nr = pp.add_run(prepared_for)
            set_font(nr, IBM, 13, "111111", bold=True)

    # Author
    author = meta.get("author")
    if author:
        ap = right.add_paragraph()
        ap.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ap.paragraph_format.space_before = Pt(0)
        ap.paragraph_format.space_after = Pt(18)
        ap.paragraph_format.line_spacing = 1.15
        al = ap.add_run("by:")
        set_font(al, EUROSTILE, 10, GRAY)
        ap.add_run().add_break()
        ar = ap.add_run(author)
        set_font(ar, IBM, 11, "111111")

    # Date only
    date_str = meta.get("date")
    if date_str:
        dp = right.add_paragraph()
        dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        dp.paragraph_format.space_before = Pt(0)
        dp.paragraph_format.space_after = Pt(0)
        dp.paragraph_format.line_spacing = 1.2
        dr = dp.add_run(date_str)
        set_font(dr, IBM, 11, "111111")


# ═══════════════════════════════════════════════════════════════════════════════
# Table of contents (standard Word TOC field with branded styles)
# ═══════════════════════════════════════════════════════════════════════════════
def setup_toc_styles(doc) -> None:
    """Configure TOC 1 / TOC 2 / TOC 3 styles with brand typography.

    Clean, modern look: no dot leaders, generous grouping around H1 entries,
    compact within groups. Page numbers right-aligned in brand blue.
    """
    tab_pos = int(7.0 * 1440)  # common right edge for TOC page numbers

    for level, font, size, color, bold, indent, sb in [
        (1, EUROSTILE_BOLD, 10.5, DARK,  True,  1.0,  10),
        (2, EUROSTILE,      9.5,  MID,   False, 1.35, 2),
        (3, EUROSTILE,      9,    GRAY,  False, 1.7,  1),
    ]:
        style_name = f"TOC {level}"
        try:
            style = doc.styles[style_name]
        except KeyError:
            from docx.enum.style import WD_STYLE_TYPE
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

        style.font.name = font
        style.font.size = Pt(size)
        style.font.color.rgb = hex_rgb(color)
        style.font.bold = bold
        pf = style.paragraph_format
        pf.space_before = Pt(sb)
        pf.space_after = Pt(1)
        pf.left_indent = Inches(indent)
        pf.line_spacing = 1.15

        # eastAsia font
        rpr = style._element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rpr.insert(0, rf)
        rf.set(qn("w:ascii"), font)
        rf.set(qn("w:hAnsi"), font)
        rf.set(qn("w:eastAsia"), font)

        # Right tab stop — NO leader (clean whitespace instead of dots)
        ppPr = style._element.find(qn("w:pPr"))
        if ppPr is None:
            ppPr = OxmlElement("w:pPr")
            style._element.append(ppPr)
        tabs_el = ppPr.find(qn("w:tabs"))
        if tabs_el is None:
            tabs_el = OxmlElement("w:tabs")
            ppPr.append(tabs_el)
        tab_el = OxmlElement("w:tab")
        tab_el.set(qn("w:val"), "right")
        tab_el.set(qn("w:leader"), "none")
        tab_el.set(qn("w:pos"), str(tab_pos))
        tabs_el.append(tab_el)


def add_toc(doc, header_rule_path: str | None = None) -> None:
    """Insert a 'contents' label + standard Word TOC field.

    The TOC field auto-populates with page numbers when opened in Word/Pages.
    """
    # "contents" label
    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(36)
    lp.paragraph_format.space_after = Pt(12)
    lp.paragraph_format.left_indent = Inches(0)
    lr = lp.add_run("contents")
    set_font(lr, EUROSTILE, 20, GRAY)

    # TOC field code
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    r1 = p.add_run()
    r1._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    r2 = p.add_run()
    r2._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r3 = p.add_run()
    r3._r.append(fld_sep)

    # Placeholder text (replaced when fields update)
    r4 = p.add_run("Table of contents will populate when you open in Word or Pages.")
    set_font(r4, IBM, 9, GRAY, italic=True)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r5 = p.add_run()
    r5._r.append(fld_end)


# ═══════════════════════════════════════════════════════════════════════════════
# Content helpers
# ═══════════════════════════════════════════════════════════════════════════════
def add_heading(
    doc, text: str, level: int, page_break: bool = False,
) -> None:
    display = text.lower() if level == 4 else text
    p = doc.add_paragraph(display, style=f"Heading {level}")
    if page_break:
        p.paragraph_format.page_break_before = True


def add_body(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.widow_control = True
    add_rich_text(p, text, IBM, 11, DARK)
    return p


def add_bullet(doc, text: str, level: int = 0):
    p = doc.add_paragraph()
    base = 1.0 + level * 0.25
    p.paragraph_format.left_indent = Inches(base + 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.widow_control = True
    br = p.add_run("\u2013  ")  # en-dash bullet
    set_font(br, EUROSTILE, 11, BLUE)
    add_rich_text(p, text, IBM, 11, DARK)
    return p


_num_counter = [0]


def add_numbered(doc, text: str, num: int | None = None):
    if num is not None:
        _num_counter[0] = num
    else:
        _num_counter[0] += 1
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.widow_control = True
    nr = p.add_run(f"{_num_counter[0]}.  ")
    set_font(nr, EUROSTILE, 11, BLUE)
    add_rich_text(p, text, IBM, 11, DARK)
    return p


def add_md_table(doc, rows_text: list[str]) -> None:
    """Parse markdown table rows and render a branded table."""
    parsed = []
    for line in rows_text:
        cells = [c.strip() for c in line.strip("|").split("|")]
        parsed.append(cells)
    data = [r for r in parsed if not all(set(c.strip()) <= {"-", ":"} for c in r)]
    if not data:
        return

    header, body_rows = data[0], data[1:]
    ncols = len(header)
    table_indent = Inches(1)
    content_w = int(Inches(6))
    col_ws = (
        [int(content_w * 0.35), int(content_w * 0.65)]
        if ncols == 2
        else [content_w // ncols] * ncols
    )

    tbl = doc.add_table(rows=1 + len(body_rows), cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = True
    tbl_pr = tbl._tbl.tblPr
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(int(table_indent.twips)))
    tbl_ind.set(qn("w:type"), "dxa")
    set_table_fixed_width(tbl)

    for j, ct in enumerate(header):
        cell = tbl.rows[0].cells[j]
        cell.width = Emu(col_ws[j])
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(ct)
        set_font(r, EUROSTILE_BOLD, 10, DARK, bold=True)
        set_cell_border(
            cell, "bottom",
            {"val": "single", "sz": "6", "color": BLUE, "space": "0"},
        )

    for i, rd in enumerate(body_rows):
        for j, ct in enumerate(rd[:ncols]):
            cell = tbl.rows[i + 1].cells[j]
            cell.width = Emu(col_ws[j])
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(ct)
            set_font(r, IBM, 10, DARK)

    # Clear default borders
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "nil")


# ═══════════════════════════════════════════════════════════════════════════════
# Markdown parser
# ═══════════════════════════════════════════════════════════════════════════════
def parse_markdown(md_text: str) -> list[tuple]:
    """Parse markdown into (type, ...) element tuples."""
    elements: list[tuple] = []
    lines = md_text.split("\n")
    i = 0
    para_lines: list[str] = []

    def flush():
        nonlocal para_lines
        if para_lines:
            elements.append(("para", " ".join(para_lines)))
            para_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            flush()
            i += 1
            continue

        if stripped.startswith("#"):
            flush()
            lvl = len(stripped) - len(stripped.lstrip("#"))
            elements.append(("heading", stripped[lvl:].strip(), lvl))
            i += 1
            continue

        if stripped == "---":
            flush()
            elements.append(("rule",))
            i += 1
            continue

        if re.match(r"^[\t ]*- ", line):
            flush()
            indent = 0
            tmp = line
            while tmp.startswith("\t"):
                indent += 1
                tmp = tmp[1:]
            text = tmp.strip()
            if text.startswith("- "):
                text = text[2:]
            elements.append(("bullet", text, indent))
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.+)", stripped)
        if m:
            flush()
            elements.append(("numbered", m.group(2), int(m.group(1))))
            i += 1
            continue

        if stripped.startswith("|"):
            flush()
            tbl = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl.append(lines[i].strip())
                i += 1
            elements.append(("table", tbl))
            continue

        para_lines.append(stripped)
        i += 1

    flush()
    return elements


def export_pdf_via_pages(docx_path: Path, pdf_path: Path) -> None:
    """Export a sibling PDF via Pages so the branded DOCX and PDF stay in sync."""
    script = f'''
    tell application "Pages"
        activate
        set docRef to open POSIX file "{docx_path}"
        delay 1
        export docRef to POSIX file "{pdf_path}" as PDF
        close docRef saving no
    end tell
    '''
    subprocess.run(["osascript", "-e", script], check=True)


def extract_frontmatter(elements: list[tuple]) -> tuple[str, list[str], int]:
    """Pull title, byline paragraphs, and content-start index from parsed elements.

    Consumes the first # heading, any immediately-following byline paragraphs
    (starting with "Prepared for" or "by:"), and any --- rules in that zone.
    Returns (title, bylines, content_start_index).
    """
    title = ""
    bylines: list[str] = []
    i = 0

    # Title
    if i < len(elements) and elements[i][0] == "heading" and elements[i][2] == 1:
        title = elements[i][1]
        i += 1

    # Bylines and rules
    while i < len(elements):
        el = elements[i]
        if el[0] == "para":
            t = el[1]
            if t.lower().startswith("prepared for") or t.lower().startswith("by:"):
                bylines.append(t)
                i += 1
                continue
        if el[0] == "rule":
            i += 1
            continue
        break

    return title, bylines, i


# ═══════════════════════════════════════════════════════════════════════════════
# Main render
# ═══════════════════════════════════════════════════════════════════════════════
def render(
    md_path: Path,
    out_path: Path,
    *,
    title: str | None = None,
    prepared_for: str | None = None,
    author: str | None = None,
    date: str | None = None,
    version: str | None = None,
    toc: str = "none",  # "none" | "after-first-section" | "after-cover"
    export_pdf: bool = True,
) -> None:
    md_text = md_path.read_text(encoding="utf-8")
    elements = parse_markdown(md_text)

    # --- Extract frontmatter from markdown ---
    md_title, md_bylines, content_start = extract_frontmatter(elements)

    # CLI overrides markdown-derived values; fall back to markdown
    final_title = title or md_title
    final_prepared_for = prepared_for
    final_author = author
    if not final_prepared_for and not final_author:
        for line in md_bylines:
            lo = line.lower()
            if lo.startswith("prepared for") and not final_prepared_for:
                # Strip leading "Prepared for " to get just the name/role
                final_prepared_for = re.sub(r"^[Pp]repared for\s+", "", line).rstrip(".")
            elif lo.startswith("by:") and not final_author:
                final_author = line[3:].strip()

    # (No pre-scan needed — standard TOC field uses heading styles directly)

    # --- Watermark ---
    wm_path = out_path.parent / ".watermark-tmp.png"
    watermark = create_footer_watermark(LOGO_PATH, wm_path)
    header_rule_path = out_path.parent / ".header-rule-tmp.png"
    header_rule = create_header_rule_image(header_rule_path)

    # --- Build document ---
    _bm_id_counter[0] = 0  # reset bookmark IDs
    _num_counter[0] = 0

    doc = Document()
    setup_page(doc, watermark_logo=watermark)
    setup_styles(doc)
    if toc != "none":
        setup_toc_styles(doc)

    # Branded header
    meta = {
        "title": final_title,
        "prepared_for": final_prepared_for,
        "author": final_author,
        "date": date,
        "version": version,
        "header_rule_path": str(header_rule) if header_rule else None,
    }
    write_branded_header(doc, meta)

    # --- TOC: "after-cover" placement ---
    toc_inserted = False
    if toc == "after-cover":
        pb = doc.add_paragraph()
        pb.add_run().add_break(WD_BREAK.PAGE)
        add_toc(doc, header_rule_path=meta.get("header_rule_path"))
        toc_inserted = True

    # --- Render body ---
    h1_count = 0

    body_elements = elements[content_start:]

    def is_listish(etype: str) -> bool:
        return etype in {"bullet", "numbered", "table"}

    for idx, el in enumerate(body_elements):
        etype = el[0]
        next_el = body_elements[idx + 1] if idx + 1 < len(body_elements) else None
        prev_el = body_elements[idx - 1] if idx > 0 else None

        if etype == "heading":
            text, md_level = el[1], el[2]
            if md_level == 1:
                continue  # additional # headings ignored (title already consumed)

            doc_level = md_level - 1  # ## → 1, ### → 2, #### → 3

            # TOC: "after-first-section" — insert before the second H1
            if doc_level == 1:
                h1_count += 1
                if toc == "after-first-section" and h1_count == 2 and not toc_inserted:
                    pb = doc.add_paragraph()
                    pb.add_run().add_break(WD_BREAK.PAGE)
                    add_toc(doc, header_rule_path=meta.get("header_rule_path"))
                    toc_inserted = True

            # Page break before H1 (except the very first, unless TOC was just placed)
            page_break = False
            if doc_level == 1:
                if h1_count > 1:
                    page_break = True
                elif h1_count == 1 and toc_inserted:
                    page_break = True  # first H1 after "after-cover" TOC

            add_heading(doc, text, doc_level, page_break=page_break)

        elif etype == "para":
            p = add_body(doc, el[1])
            text = el[1]
            is_short_para = len(text) <= 140
            follows_subhead = (
                prev_el is not None
                and prev_el[0] == "heading"
                and prev_el[2] >= 3
            )
            introduces_list = next_el is not None and is_listish(next_el[0])
            if is_short_para and (follows_subhead or introduces_list):
                p.paragraph_format.keep_together = True
            if introduces_list:
                p.paragraph_format.keep_with_next = True

        elif etype == "bullet":
            p = add_bullet(doc, el[1], level=el[2])
            starts_bullet_run = prev_el is None or prev_el[0] != "bullet"
            if starts_bullet_run and next_el is not None and next_el[0] == "bullet":
                p.paragraph_format.keep_with_next = True

        elif etype == "numbered":
            p = add_numbered(doc, el[1], num=el[2])
            starts_numbered_run = prev_el is None or prev_el[0] != "numbered"
            if starts_numbered_run and next_el is not None and next_el[0] == "numbered":
                p.paragraph_format.keep_with_next = True

        elif etype == "table":
            add_md_table(doc, el[1])

        elif etype == "rule":
            pass  # section breaks handled by H1 page breaks

    # --- Save ---
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Wrote {out_path}")
    if export_pdf:
        pdf_path = out_path.with_suffix(".pdf")
        export_pdf_via_pages(out_path.resolve(), pdf_path.resolve())
        print(f"Wrote {pdf_path}")

    # Clean up temp watermark
    if wm_path.exists():
        try:
            wm_path.unlink()
        except OSError:
            pass
    if header_rule_path.exists():
        try:
            header_rule_path.unlink()
        except OSError:
            pass


# ═══════════════════════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════════════════════
def cli() -> None:
    parser = argparse.ArgumentParser(
        description="Strike Valley Studio — Branded DOCX from Markdown",
    )
    parser.add_argument("md", type=Path, help="Path to input markdown file")
    parser.add_argument("-o", "--output", type=Path, required=True, help="Output .docx path")
    parser.add_argument("--title", help="Document title (overrides # heading)")
    parser.add_argument("--prepared-for", dest="prepared_for", help="Recipient name and role")
    parser.add_argument("--author", help="Author name and affiliation")
    parser.add_argument("--date", help="Document date")
    parser.add_argument("--version", help="Version string")
    parser.add_argument(
        "--toc",
        choices=["none", "after-first-section", "after-cover"],
        default="none",
        help="Table of contents placement",
    )
    parser.add_argument(
        "--no-pdf",
        action="store_true",
        help="Skip automatic PDF export via Pages",
    )
    args = parser.parse_args()

    render(
        args.md,
        args.output,
        title=args.title,
        prepared_for=args.prepared_for,
        author=args.author,
        date=args.date,
        version=args.version,
        toc=args.toc,
        export_pdf=not args.no_pdf,
    )


if __name__ == "__main__":
    cli()
