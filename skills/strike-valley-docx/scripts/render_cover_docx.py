"""
Purpose: Render a branded Strike Valley Studio first page into an editable DOCX.
Key exports: main(), render_document()
Role: Turn a small YAML spec into a consultancy-style cover/top-matter layout.
Invariants:
- The first page uses a 2-column table for the brand block (left) and title block (right).
- Summary sections (goal, context, methodology, etc.) are single-column paragraphs below the table.
- The output remains a normal DOCX that can be edited in Pages or Word.
"""

from __future__ import annotations

import argparse
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt, RGBColor


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Render a branded Strike Valley Studio cover/first page DOCX from a YAML spec."
    )
    parser.add_argument("--spec", required=True, type=Path, help="Path to YAML spec file.")
    parser.add_argument("--output", required=True, type=Path, help="Path to output DOCX.")
    return parser.parse_args()


def load_spec(spec_path: Path) -> dict[str, Any]:
    with spec_path.open("r", encoding="utf-8") as handle:
        spec = yaml.safe_load(handle) or {}
    if not isinstance(spec, dict):
        raise ValueError("Spec root must be a mapping.")
    return spec


def hex_to_rgb(value: str) -> RGBColor:
    normalized = value.strip().lstrip("#")
    if len(normalized) != 6:
        raise ValueError(f"Expected a 6-digit hex value, got {value!r}")
    return RGBColor.from_string(normalized.upper())


def set_run_font(run, font_name: str, size_pt: float, color_hex: str, *, bold: bool = False) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    run.font.color.rgb = hex_to_rgb(color_hex)
    run.bold = bold


def add_text(
    paragraph,
    text: str,
    *,
    font_name: str,
    size_pt: float,
    color_hex: str,
    bold: bool = False,
) -> None:
    run = paragraph.add_run(text)
    set_run_font(run, font_name, size_pt, color_hex, bold=bold)


def add_hyperlink(
    paragraph,
    *,
    text: str,
    url: str,
    font_name: str,
    size_pt: float,
    color_hex: str,
) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)
    fonts.set(qn("w:eastAsia"), font_name)
    run_properties.append(fonts)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), color_hex.strip().lstrip("#").upper())
    run_properties.append(color)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(int(size_pt * 2)))
    run_properties.append(size)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_page_background(document: Document, color_hex: str) -> None:
    background = OxmlElement("w:background")
    background.set(qn("w:color"), color_hex.strip().lstrip("#").upper())
    document._element.insert(0, background)


def clear_table_borders(table) -> None:
    table_properties = table._tbl.tblPr
    borders = table_properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "nil")


def set_cell_borders(cell, *, top: dict[str, str] | None = None) -> None:
    tc_properties = cell._tc.get_or_add_tcPr()
    tc_borders = tc_properties.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_properties.append(tc_borders)

    if top:
        top_element = tc_borders.find(qn("w:top"))
        if top_element is None:
            top_element = OxmlElement("w:top")
            tc_borders.append(top_element)
        for key, value in top.items():
            top_element.set(qn(f"w:{key}"), value)


def set_cell_margins(cell, *, top: int, start: int, bottom: int, end: int) -> None:
    tc_properties = cell._tc.get_or_add_tcPr()
    tc_mar = tc_properties.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_properties.append(tc_mar)

    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_row_height(row, height_in: float) -> None:
    row.height = Inches(height_in)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def add_page_number(paragraph, *, font_name: str, size_pt: float, color_hex: str) -> None:
    """Insert a PAGE field as properly structured separate runs, each with font properties.

    Each field component (begin, instrText, separate, display text, end) must live in its
    own <w:r> with a full <w:rPr>, otherwise Word/Pages picks up the paragraph's default
    font (often Cambria) instead of the intended one.
    """
    col = color_hex.strip().lstrip("#").upper()
    sz_val = str(int(size_pt * 2))

    def make_rpr() -> OxmlElement:
        rpr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), font_name)
        fonts.set(qn("w:hAnsi"), font_name)
        fonts.set(qn("w:eastAsia"), font_name)
        rpr.append(fonts)
        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), col)
        rpr.append(color_el)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), sz_val)
        rpr.append(sz)
        return rpr

    def run_with(*children) -> OxmlElement:
        r = OxmlElement("w:r")
        r.append(make_rpr())
        for child in children:
            r.append(child)
        return r

    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    sep   = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    text  = OxmlElement("w:t");       text.text = "1"
    end   = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")

    p = paragraph._p
    p.append(run_with(begin))
    p.append(run_with(instr))
    p.append(run_with(sep))
    p.append(run_with(text))
    p.append(run_with(end))


def configure_document(document: Document, spec: dict[str, Any]) -> None:
    style = spec.get("style", {})
    section = document.sections[0]
    section.page_width = Inches(style.get("page_width_in", 8.5))
    section.page_height = Inches(style.get("page_height_in", 11))
    section.top_margin = Inches(style.get("top_margin_in", 0.45))
    section.bottom_margin = Inches(style.get("bottom_margin_in", 0.45))
    section.left_margin = Inches(style.get("left_margin_in", 0.55))
    section.right_margin = Inches(style.get("right_margin_in", 0.55))
    section.header_distance = Inches(style.get("header_distance_in", 0.25))
    section.footer_distance = Inches(style.get("footer_distance_in", 0.25))

    body_font = style.get("body_font", "IBMPlexSans-Regular")
    text_color = style.get("text_color", "111111")
    document.styles["Normal"].font.name = body_font
    document.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), body_font)
    document.styles["Normal"].font.size = Pt(style.get("body_size_pt", 12))
    document.styles["Normal"].font.color.rgb = hex_to_rgb(text_color)

    background = style.get("page_background")
    if background:
        set_page_background(document, background)

    if style.get("page_numbers", True):
        footer_paragraph = section.footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_page_number(
            footer_paragraph,
            font_name=style.get("label_font", "Eurostile"),
            size_pt=style.get("page_number_size_pt", 9),
            color_hex=style.get("muted_color", "8A8A8A"),
        )


def write_brand_block(cell, spec: dict[str, Any], base_dir: Path) -> None:
    branding = spec.get("branding", {})
    style = spec.get("style", {})
    body_font = style.get("body_font", "IBMPlexSans-Regular")
    studio_font = style.get("studio_font", "Eurostile Bold")
    mono_font = style.get("mono_font", body_font)
    text_color = style.get("text_color", "111111")
    muted_color = style.get("muted_color", "8A8A8A")
    link_color = style.get("link_color", "1C3E8C")

    logo_path = branding.get("logo")
    if logo_path:
        resolved_logo = (base_dir / logo_path).resolve() if not Path(logo_path).is_absolute() else Path(logo_path)
        logo_paragraph = cell.paragraphs[0]
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(str(resolved_logo), width=Inches(0.85))
        logo_paragraph.paragraph_format.space_after = Pt(6)

    studio_paragraph = cell.add_paragraph()
    studio_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    studio_paragraph.paragraph_format.space_after = Pt(10)
    studio_paragraph.paragraph_format.line_spacing = 1.0
    for index, line in enumerate(branding.get("studio_name_lines", ["STRIKE VALLEY", "STUDIO"])):
        if index:
            studio_paragraph.add_run().add_break()
        add_text(
            studio_paragraph,
            line,
            font_name=studio_font,
            size_pt=12,
            color_hex=text_color,
        )

    website = branding.get("website_text")
    website_url = branding.get("website_url")
    if website and website_url:
        website_paragraph = cell.add_paragraph()
        website_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        website_paragraph.paragraph_format.space_after = Pt(1)
        website_paragraph.paragraph_format.space_before = Pt(0)
        add_hyperlink(
            website_paragraph,
            text=website,
            url=website_url,
            font_name=mono_font,
            size_pt=9,
            color_hex=link_color,
        )

    location = branding.get("location")
    if location:
        location_paragraph = cell.add_paragraph()
        location_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        location_paragraph.paragraph_format.space_before = Pt(0)
        add_text(
            location_paragraph,
            location,
            font_name=body_font,
            size_pt=9,
            color_hex=muted_color,
        )


def write_title_block(cell, spec: dict[str, Any]) -> None:
    document_spec = spec.get("document", {})
    style = spec.get("style", {})
    title_font = style.get("title_font", "EurostileRegular")
    body_font = style.get("body_font", "IBMPlexSans-Regular")
    label_font = style.get("label_font", "EurostileRegular")
    text_color = style.get("text_color", "111111")
    muted_color = style.get("muted_color", "8A8A8A")

    title_paragraph = cell.paragraphs[0]
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_paragraph.paragraph_format.space_before = Pt(6)
    title_paragraph.paragraph_format.space_after = Pt(40)
    title_paragraph.paragraph_format.line_spacing = 1.05
    lines = str(document_spec.get("title", "")).splitlines() or [""]
    for index, line in enumerate(lines):
        if index:
            title_paragraph.add_run().add_break()
        add_text(
            title_paragraph,
            line,
            font_name=title_font,
            size_pt=26,
            color_hex=text_color,
        )

    prepared_lines = [
        ("prepared_label", muted_color, False, 10),
        ("prepared_name", text_color, True, 13),
        ("prepared_role", text_color, False, 11),
    ]
    prepared_paragraph = cell.add_paragraph()
    prepared_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prepared_paragraph.paragraph_format.space_before = Pt(0)
    prepared_paragraph.paragraph_format.space_after = Pt(18)
    prepared_paragraph.paragraph_format.line_spacing = 1.15
    for index, (key, color, bold, size_pt) in enumerate(prepared_lines):
        value = document_spec.get(key)
        if not value:
            continue
        if index:
            prepared_paragraph.add_run().add_break()
        add_text(
            prepared_paragraph,
            str(value),
            font_name=label_font if key == "prepared_label" else body_font,
            size_pt=size_pt,
            color_hex=color,
            bold=bold,
        )

    date_paragraph = cell.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_paragraph.paragraph_format.space_before = Pt(0)
    date_paragraph.paragraph_format.space_after = Pt(0)
    date_paragraph.paragraph_format.line_spacing = 1.2
    for index, (key, color, size_pt) in enumerate(
        [("date", text_color, 11), ("version", muted_color, 10)]
    ):
        value = document_spec.get(key)
        if not value:
            continue
        if index:
            date_paragraph.add_run().add_break()
        add_text(
            date_paragraph,
            str(value),
            font_name=body_font,
            size_pt=size_pt,
            color_hex=color,
        )


def write_summary_section(document, row_spec: dict[str, Any], style: dict[str, Any]) -> None:
    """Write a single-column summary section (label, optional heading, body paragraphs)."""
    body_font = style.get("body_font", "IBM Plex Sans")
    label_font = style.get("label_font", "Eurostile")
    text_color = style.get("text_color", "111111")
    muted_color = style.get("muted_color", "8A8A8A")
    accent_color = row_spec.get("heading_color", style.get("accent_color", "1A5D8F"))

    # Section label (e.g. "goal", "context") — gray Eurostile lowercase
    label_paragraph = document.add_paragraph()
    label_paragraph.paragraph_format.space_before = Pt(18)
    label_paragraph.paragraph_format.space_after = Pt(4)
    add_text(
        label_paragraph,
        str(row_spec.get("label", "")).lower(),
        font_name=label_font,
        size_pt=16,
        color_hex=muted_color,
    )

    # Indent for body content (labels stay flush left)
    body_indent = Inches(1)

    # Optional heading (e.g. "Discovery interviews (8-10 hours)")
    heading = row_spec.get("heading")
    if heading:
        heading_paragraph = document.add_paragraph()
        heading_paragraph.paragraph_format.space_before = Pt(0)
        heading_paragraph.paragraph_format.space_after = Pt(6)
        heading_paragraph.paragraph_format.line_spacing = 1.1
        heading_paragraph.paragraph_format.left_indent = body_indent
        add_text(
            heading_paragraph,
            str(heading),
            font_name=label_font if row_spec.get("heading_font") == "title" else body_font,
            size_pt=row_spec.get("heading_size_pt", 13),
            color_hex=accent_color,
            bold=True,
        )

    # Body paragraphs
    body_size = row_spec.get("body_size_pt", style.get("body_size_pt", 11))
    paragraphs_list = row_spec.get("paragraphs", [])
    for index, paragraph_text in enumerate(paragraphs_list):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.2
        paragraph.paragraph_format.space_after = Pt(6 if index < len(paragraphs_list) - 1 else 2)
        paragraph.paragraph_format.left_indent = body_indent
        add_text(
            paragraph,
            str(paragraph_text),
            font_name=body_font,
            size_pt=body_size,
            color_hex=text_color,
        )


def render_document(spec: dict[str, Any], base_dir: Path) -> Document:
    style = spec.get("style", {})
    layout = spec.get("layout", {})

    document = Document()
    configure_document(document, spec)

    grid = document.add_table(rows=1, cols=2)
    grid.alignment = WD_TABLE_ALIGNMENT.LEFT
    grid.autofit = False
    clear_table_borders(grid)

    left_width_in = layout.get("left_column_in", 2.2)
    right_width_in = layout.get("right_column_in", 5.1)
    grid.columns[0].width = Inches(left_width_in)
    grid.columns[1].width = Inches(right_width_in)

    top_row = grid.rows[0]
    top_row.cells[0].width = Inches(left_width_in)
    top_row.cells[1].width = Inches(right_width_in)
    top_row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    top_row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_row_height(top_row, layout.get("top_row_min_height_in", 2.2))
    set_cell_margins(top_row.cells[0], top=50, start=24, bottom=120, end=50)
    set_cell_margins(top_row.cells[1], top=80, start=110, bottom=120, end=10)
    set_cell_borders(
        top_row.cells[1],
        top={"val": "single", "sz": "8", "space": "4", "color": "222222"},
    )

    write_brand_block(top_row.cells[0], spec, base_dir)
    write_title_block(top_row.cells[1], spec)

    # Summary sections as single-column paragraphs below the table
    for row_spec in spec.get("cover_rows", []):
        write_summary_section(document, row_spec, style)

    if spec.get("body", {}).get("page_break_after_cover"):
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    return document


def main() -> None:
    args = parse_args()
    spec = load_spec(args.spec)
    document = render_document(spec, args.spec.parent)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)
    print(f"Wrote {args.output}")


if __name__ == "__main__":
    main()
